#!/usr/bin/env python3
"""Build local IEEE journal and word-abbreviation rule databases.

The builder is deliberately offline.  It accepts the locally supplied IEEE
guide as DOCX or PDF and the IEEE journal-title workbook as XLSX.  Every
generated rule retains an auditable source location.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable


def fail_dependency(package: str, exc: Exception) -> None:
    print(
        f"Missing dependency '{package}': {exc}\n"
        "Install the skill dependencies with:\n"
        "  python -m pip install -r requirements.txt",
        file=sys.stderr,
    )
    raise SystemExit(2)


try:
    import openpyxl
except ImportError as exc:  # pragma: no cover - exercised only without deps
    fail_dependency("openpyxl", exc)


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_GUIDE_DOCX = ROOT / "references" / "IEEE_Reference_Style_Guide_for_Authors.docx"
DEFAULT_GUIDE_PDF = ROOT / "references" / "IEEE_Reference_Style_Guide_for_Authors.pdf"
DEFAULT_XLSX = (
    ROOT
    / "references"
    / "List_of_IEEE_Journal_Magazine_Titles_Internal_Acronym_and_Reference_Abbreviation.xlsx"
)
DATA_DIR = ROOT / "data"


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def norm_header(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", " ", str(value or "").casefold()).strip()


def compact_space(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def canonical_full_title(title: str) -> str:
    """Convert workbook inverted titles to their ordinary reading order."""
    title = compact_space(title)
    patterns = (
        r"^(?P<body>.+),\s*(?P<prefix>IEEE/ACM Transactions on)$",
        r"^(?P<body>.+),\s*(?P<prefix>IEEE Transactions on)$",
        r"^(?P<body>.+),\s*(?P<prefix>IEEE Journal of)$",
        r"^(?P<body>.+),\s*(?P<prefix>IEEE Letters on)$",
    )
    for pattern in patterns:
        match = re.match(pattern, title, flags=re.IGNORECASE)
        if match:
            return f"{match.group('prefix')} {match.group('body')}"
    match = re.match(r"^(?P<body>.+),\s*(?P<prefix>IEEE)$", title, re.IGNORECASE)
    if match:
        return f"IEEE {match.group('body')}"
    return title


def split_lines(value: Any) -> list[str]:
    return [compact_space(part) for part in str(value or "").splitlines() if compact_space(part)]


def primary_acronym(value: Any) -> str:
    lines = split_lines(value)
    for line in lines:
        candidate = line.replace("*", "").strip()
        if candidate:
            return candidate
    return ""


def primary_reference_abbreviation(value: Any) -> str:
    lines = split_lines(value)
    for line in lines:
        if not line:
            continue
        candidate = line.replace("*", "").strip()
        candidate = re.sub(
            r"\s*\((?:19|20)\d{2}(?:\s*[-–]\s*(?:present|(?:19|20)\d{2}))?\)\s*$",
            "",
            candidate,
            flags=re.IGNORECASE,
        ).strip()
        candidate = re.sub(
            r"\s*\((?:from\s+)?(?:Jan\.?|Feb\.?|Mar\.?|Apr\.?|May|Jun\.?|Jul\.?|Aug\.?|Sep\.?|Sept\.?|Oct\.?|Nov\.?|Dec\.?)?\s*(?:19|20)\d{2}[^)]*\)\s*$",
            "",
            candidate,
            flags=re.IGNORECASE,
        ).strip()
        if candidate:
            return candidate
    return ""


def locate_header(ws: Any) -> tuple[int, dict[str, int]]:
    """Find semantic columns without assuming a fixed row or column."""
    best: tuple[int, dict[str, int]] | None = None
    for row_no in range(1, min(ws.max_row or 1, 25) + 1):
        values = [ws.cell(row_no, col).value for col in range(1, (ws.max_column or 1) + 1)]
        headers = {col + 1: norm_header(value) for col, value in enumerate(values) if value is not None}
        ref_cols = [c for c, h in headers.items() if "reference" in h and "abbreviation" in h]
        if not ref_cols:
            continue
        full_cols = [
            c
            for c, h in headers.items()
            if h in {"title", "full title", "journal title", "journal magazine title"}
            or ("full" in h and "title" in h)
        ]
        acronym_cols = [c for c, h in headers.items() if "internal" in h and "acronym" in h]
        if not acronym_cols:
            acronym_cols = [
                c
                for c, h in headers.items()
                if h in {"journal magazine", "acronym", "journal magazine acronym"}
            ]
        if full_cols:
            mapping = {
                "full_title": full_cols[0],
                "reference_abbreviation": ref_cols[0],
            }
            if acronym_cols:
                mapping["internal_acronym"] = acronym_cols[0]
            best = (row_no, mapping)
            break
    if best is None:
        raise ValueError(
            "Could not identify XLSX headers. Expected semantic columns such as "
            "Title/Full Title, Internal Acronym or Journal/Magazine, and Reference Abbreviation."
        )
    return best


def build_journal_rules(path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    try:
        # Some workbooks omit cached dimension metadata, causing read-only
        # worksheets to report max_row/max_column as None.  This official list
        # is small, so normal mode is safer and still inexpensive.
        workbook = openpyxl.load_workbook(path, read_only=False, data_only=True)
    except Exception as exc:
        raise RuntimeError(f"Unable to read XLSX '{path}': {exc}") from exc

    records: list[dict[str, Any]] = []
    unresolved: list[dict[str, Any]] = []
    sheets_meta: list[dict[str, Any]] = []
    for ws in workbook.worksheets:
        try:
            header_row, columns = locate_header(ws)
        except ValueError as exc:
            unresolved.append({"sheet": ws.title, "row": None, "reason": str(exc)})
            continue
        sheets_meta.append({"sheet": ws.title, "header_row": header_row, "columns": columns})
        for row_no in range(header_row + 1, (ws.max_row or header_row) + 1):
            raw_full = ws.cell(row_no, columns["full_title"]).value
            raw_ref = ws.cell(row_no, columns["reference_abbreviation"]).value
            raw_acronym = (
                ws.cell(row_no, columns["internal_acronym"]).value
                if "internal_acronym" in columns
                else None
            )
            if not any((raw_full, raw_ref, raw_acronym)):
                continue
            full = canonical_full_title(str(raw_full or ""))
            ref = primary_reference_abbreviation(raw_ref)
            acronym = primary_acronym(raw_acronym)
            if not full or not ref:
                unresolved.append(
                    {
                        "sheet": ws.title,
                        "row": row_no,
                        "full_title": raw_full,
                        "internal_acronym": raw_acronym,
                        "reference_abbreviation": raw_ref,
                        "reason": "missing usable full title or current reference abbreviation",
                    }
                )
                continue
            aliases = []
            for alias in [str(raw_full or ""), full, acronym, ref, *split_lines(raw_ref), *split_lines(raw_acronym)]:
                alias = compact_space(alias.replace("*", ""))
                if alias and alias not in aliases:
                    aliases.append(alias)
            records.append(
                {
                    "full_title": full,
                    "workbook_title": compact_space(raw_full),
                    "internal_acronym": acronym,
                    "reference_abbreviation": ref,
                    "aliases": aliases,
                    "source": {
                        "source_file": path.name,
                        "sheet": ws.title,
                        "row": row_no,
                        "match_columns": columns,
                    },
                }
            )
    return records, unresolved, {"sheets": sheets_meta}


def add_word_rule(
    rules: dict[str, dict[str, Any]],
    word: str,
    abbreviation: str,
    source_file: str,
    section: str,
    table: int | None,
    row: int | None,
) -> None:
    word = compact_space(word)
    abbreviation = compact_space(abbreviation)
    if not word or not abbreviation or len(word) == 1:
        return
    key = word.casefold()
    candidate = {
        "word": word,
        "abbreviation": abbreviation,
        "source": {
            "source_file": source_file,
            "section": section,
            "table": table,
            "row": row,
        },
    }
    if key not in rules:
        rules[key] = candidate


def build_word_rules_docx(path: Path) -> tuple[dict[str, dict[str, Any]], list[str], dict[str, Any]]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        fail_dependency("python-docx", exc)
    try:
        document = Document(path)
    except Exception as exc:
        raise RuntimeError(f"Unable to read DOCX guide '{path}': {exc}") from exc

    rules: dict[str, dict[str, Any]] = {}
    unresolved: list[str] = []
    common_heading = "Common Abbreviations of Words in References"
    heading_found = any(common_heading.casefold() in p.text.casefold() for p in document.paragraphs)
    if not heading_found:
        unresolved.append(f"Section heading not found: {common_heading}")

    # The supplied guide stores conference terms in one 4-column table and the
    # common word list in a two-column table.  Detect by shape/content rather
    # than fixed document table numbers.
    detected_tables: list[dict[str, Any]] = []
    for table_no, table in enumerate(document.tables, start=1):
        rows = [[compact_space(cell.text) for cell in row.cells] for row in table.rows]
        flat = " ".join(value for row in rows[:20] for value in row).casefold()
        if "conference" in flat and "conf." in flat and "symposium" in flat:
            section = "Conferences and Conference Proceedings / common conference abbreviations"
            pairs = [(0, 1), (2, 3)]
        elif len(rows) >= 100 and any(row and row[0].casefold() == "abstracts" for row in rows[:20]):
            section = common_heading
            pairs = [(0, 1)]
        else:
            continue
        detected_tables.append({"table": table_no, "rows": len(rows), "section": section})
        for row_no, row in enumerate(rows, start=1):
            for left, right in pairs:
                if len(row) <= right:
                    continue
                add_word_rule(rules, row[left], row[right], path.name, section, table_no, row_no)

    if not detected_tables:
        unresolved.append("No reliable abbreviation tables detected in DOCX guide.")
    return rules, unresolved, {"detected_tables": detected_tables, "paragraphs": len(document.paragraphs)}


def iter_pdf_pages(path: Path) -> Iterable[tuple[int, str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        fail_dependency("pypdf", exc)
    try:
        reader = PdfReader(path)
        for index, page in enumerate(reader.pages, start=1):
            yield index, page.extract_text() or ""
    except Exception as exc:
        raise RuntimeError(f"Unable to extract PDF text layer from '{path}': {exc}") from exc


def build_word_rules_pdf(path: Path) -> tuple[dict[str, dict[str, Any]], list[str], dict[str, Any]]:
    pages = list(iter_pdf_pages(path))
    if not any(text.strip() for _, text in pages):
        raise RuntimeError(
            f"PDF '{path}' has no usable text layer. OCR is not run automatically; provide a text-layer PDF or DOCX."
        )
    rules: dict[str, dict[str, Any]] = {}
    unresolved: list[str] = []
    section_pages = [
        number
        for number, text in pages
        if "Common Abbreviations of Words in References".casefold() in text.casefold()
    ]
    if not section_pages:
        unresolved.append("Could not locate 'Common Abbreviations of Words in References' in PDF text.")
        return rules, unresolved, {"pages": len(pages), "section_pages": []}
    start = section_pages[0]
    for page_no, text in pages:
        if page_no < start or page_no > start + 8:
            continue
        for line_no, line in enumerate(text.splitlines(), start=1):
            line = compact_space(line)
            match = re.match(r"^([A-Za-z][A-Za-z -]{2,}?)\s{2,}([A-Za-z][A-Za-z.-]{1,})$", line)
            if match:
                add_word_rule(
                    rules,
                    match.group(1),
                    match.group(2),
                    path.name,
                    "Common Abbreviations of Words in References",
                    None,
                    page_no,
                )
            elif line and len(line) < 100:
                unresolved.append(f"page {page_no}, line {line_no}: {line}")
    return rules, unresolved, {"pages": len(pages), "section_pages": section_pages}


def write_json(path: Path, data: Any) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def choose_guide(path_arg: str | None) -> Path:
    if path_arg:
        return Path(path_arg).resolve()
    if DEFAULT_GUIDE_DOCX.exists():
        return DEFAULT_GUIDE_DOCX
    return DEFAULT_GUIDE_PDF


def build(guide: Path, journals: Path, output_dir: Path) -> dict[str, int]:
    if not guide.exists():
        raise FileNotFoundError(f"IEEE guide not found: {guide}")
    if not journals.exists():
        raise FileNotFoundError(f"IEEE journal workbook not found: {journals}")
    output_dir.mkdir(parents=True, exist_ok=True)

    journal_records, unresolved_xlsx, xlsx_meta = build_journal_rules(journals)
    if guide.suffix.casefold() == ".docx":
        word_rules, unresolved_guide, guide_meta = build_word_rules_docx(guide)
        extraction_method = "python-docx table extraction"
    elif guide.suffix.casefold() == ".pdf":
        word_rules, unresolved_guide, guide_meta = build_word_rules_pdf(guide)
        extraction_method = "pypdf text-layer extraction"
    else:
        raise ValueError("--guide must be a .docx or .pdf file")

    generated_at = utc_now()
    journals_payload = {
        "metadata": {
            "generated_at": generated_at,
            "source_file": journals.name,
            "record_count": len(journal_records),
            **xlsx_meta,
        },
        "records": journal_records,
    }
    words_payload = {
        "metadata": {
            "generated_at": generated_at,
            "source_file": guide.name,
            "extraction_method": extraction_method,
            "record_count": len(word_rules),
            **guide_meta,
        },
        "abbreviations": dict(sorted(word_rules.items())),
    }
    sources_payload = {
        "metadata": {"generated_at": generated_at},
        "rules": {
            "IEEE-JOURNAL-001": {
                "source_file": journals.name,
                "section": "Journal/Magazine title table",
                "location": "per-record worksheet row",
            },
            "IEEE-CONF-001": {
                "source_file": guide.name,
                "section": "Conferences and Conference Proceedings",
                "keywords": ["Proceedings", "Conference", "Symposium", "International", "Annual"],
            },
            "IEEE-VENUE-001": {
                "source_file": guide.name,
                "section": "Common Abbreviations of Words in References",
                "location": "per-word table row or PDF page",
            },
            "IEEE-AUTHOR-001": {
                "source_file": guide.name,
                "section": "II. Style",
                "keywords": ["list names of all authors", "et al."],
            },
            "IEEE-DOI-001": {
                "source_file": guide.name,
                "section": "Periodicals / Conference Proceedings With DOI",
                "keywords": ["doi"],
            },
            "IEEE-PAGES-001": {
                "source_file": guide.name,
                "section": "Periodicals / Conference Proceedings",
                "keywords": ["pp.", "Art no."],
            },
            "IEEE-TITLE-001": {
                "source_file": "SKILL.md",
                "section": "Safe edit policy / title case protection",
                "basis": "BibTeX case-changing styles require braces around confirmed acronyms and system names.",
            },
            "IEEE-SYNTAX-004": {
                "source_file": "SKILL.md",
                "section": "Safe edit policy / minimal BibTeX normalization",
                "basis": "Known field names are normalized to lowercase without changing their values or order.",
            },
            "IEEE-FIELD-001": {
                "source_file": guide.name,
                "section": "Basic Format blocks by reference type",
                "keywords": ["Basic Format"],
            },
            "IEEE-CONF-003": {
                "source_file": guide.name,
                "section": "Conferences and Conference Proceedings / Common Abbreviations of Words in References",
                "keywords": ["Proc.", "Conf.", "Symp.", "Int.", "Annu."],
            },
            "IEEE-JOURNAL-010": {
                "source_file": guide.name,
                "section": "Periodicals / Common Abbreviations of Words in References",
                "policy": "Only sourced word rules are applied; unresolved words are retained for manual review.",
            },
        },
    }
    write_json(output_dir / "ieee_journal_abbreviations.json", journals_payload)
    write_json(output_dir / "ieee_word_abbreviations.json", words_payload)
    write_json(output_dir / "rule_sources.json", sources_payload)
    write_json(output_dir / "unresolved_xlsx_rows.json", unresolved_xlsx)
    unresolved_text = [
        "# Lines or sections that could not be parsed reliably",
        f"# Source: {guide.name}",
        *unresolved_guide,
    ]
    (output_dir / "unresolved_pdf_rules.txt").write_text(
        "\n".join(unresolved_text).rstrip() + "\n", encoding="utf-8"
    )
    return {
        "journals": len(journal_records),
        "words": len(word_rules),
        "unresolved_xlsx": len(unresolved_xlsx),
        "unresolved_guide": len(unresolved_guide),
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--guide", help="Local IEEE guide (.docx or text-layer .pdf)")
    parser.add_argument("--journals", default=str(DEFAULT_XLSX), help="Local IEEE journal-title XLSX")
    parser.add_argument("--output-dir", default=str(DATA_DIR), help="Directory for generated JSON data")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    guide = choose_guide(args.guide)
    try:
        counts = build(guide, Path(args.journals).resolve(), Path(args.output_dir).resolve())
    except Exception as exc:
        print(f"Rule database build failed: {exc}", file=sys.stderr)
        return 1
    print(
        "Built IEEE rule database: "
        f"{counts['journals']} journal rows, {counts['words']} word abbreviations, "
        f"{counts['unresolved_xlsx']} unresolved XLSX rows, "
        f"{counts['unresolved_guide']} unresolved guide items."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
